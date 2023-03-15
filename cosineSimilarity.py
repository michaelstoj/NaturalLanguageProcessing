import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import docx

def preprocess(text):
  # remove punctuation and make all text lowercase
  text = re.sub(r'[^\w\s]', '', text)
  text = text.lower()
  
  # split text into a list of words
  words = text.split()
  
  return words


doc1 = docx.Document("article1.docx")
text1 = []

for para in doc1.paragraphs:
    text1.append(para.text)

doc2 = docx.Document("article2.docx")
text2 = []

for para in doc2.paragraphs:
    text2.append(para.text)
    
# define the documents to compare
doc1 = " ".join(text1)
doc2 = " ".join(text2)

# preprocess the documents
doc1_words = preprocess(doc1)
doc2_words = preprocess(doc2)

# create a TfidfVectorizer object
vectorizer = TfidfVectorizer()

doc1_string = " ".join(doc1_words)
doc2_string = " ".join(doc2_words)

# generate tf-idf matrix
tfidf_matrix = vectorizer.fit_transform([doc1_string, doc2_string])

# get the cosine similarity of the two documents
cosine_similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix)

print(cosine_similarity)


textss = [doc1, doc2]
# Compute the TF-IDF vectors of the texts
vectorizer = TfidfVectorizer()
tfidf_vectors = vectorizer.fit_transform(textss)

# Compute the cosine similarity between all pairs of texts
similarity = cosine_similarity(tfidf_vectors)

print(similarity)